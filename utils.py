import os
import re
from docx import Document
from docx.oxml.ns import qn

DEFAULT_IMAGE_DIR = "static/question_images"

def extract_table_html(table):
    html = "<table border='1' cellspacing='0' cellpadding='5'>"
    for row in table.rows:
        html += "<tr>"
        for cell in row.cells:
            html += f"<td>{cell.text.strip()}</td>"
        html += "</tr>"
    html += "</table>"
    return html

def save_image_from_run(run, output_dir, image_counter):
    blip_elements = run._element.findall('.//a:blip', namespaces={
        'a': 'http://schemas.openxmlformats.org/drawing/2006/main'
    })

    if not blip_elements:
        return None

    rId = blip_elements[0].get(qn('r:embed'))
    image_part = run.part.related_parts[rId]
    image_data = image_part.blob

    image_filename = f"question_image_{image_counter}.png"
    image_path = os.path.join(output_dir, image_filename)

    with open(image_path, 'wb') as f:
        f.write(image_data)

    return image_filename

def parse_docx_questions(file_stream, image_output_dir=DEFAULT_IMAGE_DIR):
    document = Document(file_stream)
    questions = []
    current_question = None
    extra_html_parts = []
    image_counter = 0
    question_number = 1
    found_first_question = False
    skipped = 0

    os.makedirs(image_output_dir, exist_ok=True)

    for para in document.paragraphs:
        text = para.text.strip()

        # Handle embedded images
        for run in para.runs:
            image_name = save_image_from_run(run, image_output_dir, image_counter + 1)
            if image_name and current_question:
                image_counter += 1
                current_question["image"] = image_name

        if not text:
            continue

        # Detect question start (with spacing flexibility)
        if re.match(rf"^\s*{question_number}[\.\)]\s+", text):
            # Save previous question
            if current_question:
                current_question["extra_content"] = ''.join(extra_html_parts) if extra_html_parts else None
                if current_question.get("question") and current_question.get("answer") in ["a", "b", "c", "d"]:
                    questions.append(current_question)
                else:
                    skipped += 1
            extra_html_parts = []

            # Extract marks and clean question text
            marks_match = re.search(r"\((\d+)\s?(?:mks|marks?)\)", text, re.IGNORECASE)
            marks = int(marks_match.group(1)) if marks_match else 1
            clean_text = re.sub(r"\s*\(\d+\s?(?:mks|marks?)\)", "", text)

            question_text = re.sub(rf"^\s*{question_number}[\.\)]\s+", "", clean_text)
            current_question = {
                "question": question_text,
                "a": "", "b": "", "c": "", "d": "",
                "answer": "",
                "extra_content": None,
                "image": None,
                "marks": marks
            }
            question_number += 1
            found_first_question = True
            continue

        # Options A-D
        if re.match(r"^\(?[a-dA-D][\.\)]", text) and current_question:
            match = re.match(r"^\(?([a-dA-D])[\.\)]\s*(.+)", text)
            if match:
                label = match.group(1).lower()
                content = match.group(2).strip()
                current_question[label] = content
            continue

        # Answer line
        if re.match(r"^(answer|correct answer):", text, re.IGNORECASE) and current_question:
            match = re.search(r":\s*([a-dA-D])", text, re.IGNORECASE)
            if match:
                current_question["answer"] = match.group(1).lower()
            continue

        # Extra explanation under the current question
        if found_first_question:
            extra_html_parts.append(f"<p>{text}</p>")

    # Final question
    if current_question:
        current_question["extra_content"] = ''.join(extra_html_parts) if extra_html_parts else None
        if current_question.get("question") and current_question.get("answer") in ["a", "b", "c", "d"]:
            questions.append(current_question)
        else:
            skipped += 1

    # Append tables to the last question
    for table in document.tables:
        if current_question:
            table_html = extract_table_html(table)
            current_question["extra_content"] = (current_question.get("extra_content") or '') + table_html

    print(f"✅ Parsed {len(questions)} valid questions.")
    if skipped > 0:
        print(f" Skipped {skipped} question(s) due to missing options or answers.")

    return questions

def get_quiz_status(session, quiz_id, student_id):
    from models import Result
    result = session.query(Result).filter_by(quiz_id=quiz_id, student_id=student_id).first()
    return 'Completed' if result else 'Pending'